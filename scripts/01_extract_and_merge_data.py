# -*- coding: utf-8 -*-
"""
从 raw Word 文档抽取结构化景点数据和原文知识块。

默认只处理 data/raw 中的两个 docx 文档，不依赖大模型生成内容。
"""

import json
import os
import re
from collections import defaultdict
from docx import Document


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
DATA_RAW = os.path.join(PROJECT_ROOT, "data", "raw")
DATA_PROCESSED = os.path.join(PROJECT_ROOT, "data", "processed")

DOC_ATTRACTIONS = os.path.join(DATA_RAW, "灵山胜境 景点结构化数据集.docx")
DOC_GUIDE = os.path.join(DATA_RAW, "灵山胜境：历史、文化、景点特色与个性化游览指南.docx")

OUTPUT_ATTRACTIONS = os.path.join(DATA_PROCESSED, "attractions_complete.json")
OUTPUT_KNOWLEDGE_UNITS = os.path.join(DATA_PROCESSED, "knowledge_units.json")
OUTPUT_SOURCE_SUMMARY = os.path.join(DATA_PROCESSED, "source_summary.json")

ATTRACTION_COLUMNS = {
    "景区名称": "scenic_name",
    "景点ID": "id",
    "景点名称": "name",
    "具体位置": "location",
    "建筑/景观参数": "specs",
    "核心功能": "function",
    "文化内涵": "culture",
    "详细介绍": "details",
    "游玩亮点": "highlights",
    "演艺/开放信息": "hours",
    "备注": "remarks",
}

SUPPLEMENT_TABLE_NAMES = [
    "灵山大佛",
    "灵山梵宫",
    "九龙灌浴",
    "五印坛城",
    "祥符禅寺",
]

STOP_HEADINGS = {
    "数据集说明",
    "子表1：灵山胜境 景点数据集",
    "表2：拈花湾禅意小镇 景点数据集",
}


def clean_text(text: str) -> str:
    text = text or ""
    text = text.replace("\u3000", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_key(text: str) -> str:
    return clean_text(text).replace(" ", "")


def ensure_output_dir():
    os.makedirs(DATA_PROCESSED, exist_ok=True)


def require_files():
    missing = [path for path in [DOC_ATTRACTIONS, DOC_GUIDE] if not os.path.exists(path)]
    if missing:
        for path in missing:
            print(f"缺少原始文档: {path}")
        raise FileNotFoundError("raw Word 文档不完整")


def table_to_matrix(table):
    return [[clean_text(cell.text) for cell in row.cells] for row in table.rows]


def extract_attraction_tables(doc_path: str) -> list:
    doc = Document(doc_path)
    attractions = []

    for table in doc.tables:
        rows = table_to_matrix(table)
        if not rows:
            continue
        headers = rows[0]
        if "景区名称" not in headers or "景点名称" not in headers:
            continue

        header_map = {
            index: ATTRACTION_COLUMNS.get(header)
            for index, header in enumerate(headers)
            if ATTRACTION_COLUMNS.get(header)
        }

        for row in rows[1:]:
            item = {field: "" for field in ATTRACTION_COLUMNS.values()}
            for index, field in header_map.items():
                if index < len(row):
                    item[field] = row[index]

            if not item.get("name"):
                continue

            scenic_name = item.get("scenic_name", "")
            if "灵山胜境" in scenic_name:
                item["scenic_type"] = "lingshan"
            elif "拈花湾" in scenic_name:
                item["scenic_type"] = "nianhuawan"
            else:
                item["scenic_type"] = "other"

            item["aliases"] = build_aliases(item["name"])
            item["supplement"] = {}
            item["source_files"] = [os.path.basename(doc_path)]
            attractions.append(item)

    return attractions


def build_aliases(name: str) -> list:
    aliases = {name}
    if name.startswith("灵山"):
        aliases.add(name.replace("灵山", "", 1))
    if name.endswith("景区"):
        aliases.add(name[:-2])
    return sorted(alias for alias in aliases if alias)


def extract_guide_tables(doc_path: str) -> list:
    doc = Document(doc_path)
    tables = []
    for index, table in enumerate(doc.tables):
        rows = table_to_matrix(table)
        if len(rows) < 2:
            continue

        headers = rows[0]
        table_data = {
            "index": index,
            "title": infer_table_title(index, headers),
            "headers": headers,
            "rows": rows[1:],
            "items": [],
            "source_file": os.path.basename(doc_path),
        }

        if len(headers) == 2 and headers[0] in {"项目", "票种"}:
            for row in rows[1:]:
                if len(row) >= 2 and row[0] and row[1]:
                    table_data["items"].append({"key": row[0], "value": row[1]})
        elif len(headers) >= 3:
            for row in rows[1:]:
                item = {}
                for i, header in enumerate(headers):
                    item[header or f"字段{i + 1}"] = row[i] if i < len(row) else ""
                if any(item.values()):
                    table_data["items"].append(item)

        tables.append(table_data)
    return tables


def infer_table_title(index: int, headers: list) -> str:
    if index < len(SUPPLEMENT_TABLE_NAMES):
        return SUPPLEMENT_TABLE_NAMES[index]
    if headers and "票种" in headers[0]:
        return "票务价格"
    return f"补充表{index + 1}"


def merge_supplements(attractions: list, guide_tables: list) -> list:
    by_name = {item["name"]: item for item in attractions}
    merged = 0

    for table in guide_tables:
        title = table.get("title", "")
        attraction = by_name.get(title)
        if not attraction:
            continue

        for item in table.get("items", []):
            if "key" in item:
                attraction["supplement"][item["key"]] = item["value"]
        if table.get("items"):
            attraction["source_files"].append(table["source_file"])
            merged += 1

    print(f"补充表合并到景点: {merged} 张")
    return attractions


def paragraph_level(style_name: str) -> int:
    match = re.search(r"Heading\s+(\d+)|标题\s*(\d+)", style_name or "", re.I)
    if not match:
        return 0
    return int(match.group(1) or match.group(2))


def looks_like_heading(text: str) -> bool:
    if not text or len(text) > 40:
        return False
    if text in STOP_HEADINGS:
        return True
    if text.endswith(("：", "详解", "指南")):
        return True
    if re.match(r"^\d{4}年", text):
        return False
    return bool(re.search(r"(概况|历史|文化|路线|攻略|景点|票务|设施|贴士|指南|体验|推荐)", text))


def split_guide_paragraphs(doc_path: str) -> list:
    doc = Document(doc_path)
    units = []
    current_title = "文档概览"
    current_parts = []

    def flush():
        nonlocal current_parts
        content = clean_text(" ".join(current_parts))
        if len(content) >= 30:
            units.append({
                "title": current_title,
                "content": content,
                "category": classify_text(f"{current_title} {content}", title=current_title, source_type="paragraph"),
                "source_file": os.path.basename(doc_path),
                "source_type": "paragraph",
            })
        current_parts = []

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        level = paragraph_level(getattr(para.style, "name", ""))
        if level or looks_like_heading(text):
            flush()
            current_title = text
            continue

        current_parts.append(text)
        if sum(len(part) for part in current_parts) >= 700:
            flush()

    flush()
    return units


def attraction_to_knowledge_units(attractions: list) -> list:
    units = []
    field_groups = [
        ("基础信息", ["location", "specs", "function"], "景点介绍"),
        ("文化内涵", ["culture"], "历史文化"),
        ("详细介绍", ["details"], "景点介绍"),
        ("游玩亮点", ["highlights", "hours", "remarks"], "游玩建议"),
    ]
    labels = {
        "location": "位置",
        "specs": "建筑/景观参数",
        "function": "核心功能",
        "culture": "文化内涵",
        "details": "详细介绍",
        "highlights": "游玩亮点",
        "hours": "演艺/开放信息",
        "remarks": "备注",
    }

    for attraction in attractions:
        name = attraction["name"]
        for group_name, fields, category in field_groups:
            parts = []
            for field in fields:
                value = attraction.get(field)
                if value:
                    parts.append(f"{labels[field]}：{value}")
            if parts:
                units.append({
                    "title": f"{name}{group_name}",
                    "content": "\n".join(parts),
                    "category": category,
                    "source_file": "；".join(sorted(set(attraction.get("source_files", [])))),
                    "source_type": "attraction",
                    "source_attraction": attraction.get("id", ""),
                    "attraction_name": name,
                })

        for key, value in attraction.get("supplement", {}).items():
            if value:
                units.append({
                    "title": f"{name}{key}",
                    "content": f"{key}：{value}",
                    "category": classify_text(f"{name} {key} {value}", title=key, source_type="supplement_table"),
                    "source_file": "灵山胜境：历史、文化、景点特色与个性化游览指南.docx",
                    "source_type": "supplement_table",
                    "source_attraction": attraction.get("id", ""),
                    "attraction_name": name,
                })
    return units


def guide_tables_to_knowledge_units(tables: list) -> list:
    units = []
    for table in tables:
        title = table.get("title", "")
        if title in SUPPLEMENT_TABLE_NAMES:
            continue
        rows = []
        for item in table.get("items", []):
            if "key" in item:
                rows.append(f"{item['key']}：{item['value']}")
            else:
                rows.append("；".join(f"{k}：{v}" for k, v in item.items() if v))
        content = "\n".join(row for row in rows if row)
        if content:
            units.append({
                "title": title,
                "content": content,
                "category": classify_text(f"{title} {content}", title=title, source_type="table"),
                "source_file": table["source_file"],
                "source_type": "table",
            })
    return units


def classify_text(text: str, title: str = "", source_type: str = "") -> str:
    title = title or ""
    if source_type == "table":
        if any(word in title for word in ["票务", "票价", "门票", "价格"]):
            return "票务价格"
        if any(word in title for word in ["开放时间", "营业时间", "演出时间", "表演时间"]):
            return "开放时间"
        if any(word in title for word in ["交通", "路线", "怎么去"]):
            return "交通路线"

    if source_type == "supplement_table":
        if any(word in title for word in ["表演", "演出", "场次", "开放", "时间"]):
            return "开放时间"
        if any(word in title for word in ["体验", "游玩", "看点", "亮点", "最佳"]):
            return "游玩建议"
        if any(word in title for word in ["意义", "文化", "地位", "艺术", "历史", "缘起"]):
            return "历史文化"
        if any(word in title for word in ["数据", "规模", "高度", "面积", "建筑"]):
            return "景点数据"

    if any(word in title for word in ["亲子", "自然风光", "爱好者路线", "游览路线", "游玩路线", "路线"]):
        return "游玩建议"
    if any(word in title for word in ["历史", "文化", "缘起", "兴衰", "交流平台"]):
        return "历史文化"
    if any(word in title for word in ["票务", "票价", "门票", "优惠票"]):
        return "票务价格"
    if any(word in title for word in ["开放时间", "营业时间", "表演时间", "演出时间"]):
        return "开放时间"
    if any(word in title for word in ["交通", "停车", "公交", "自驾"]):
        return "交通路线"

    rules = [
        ("交通路线", ["交通", "公交", "自驾", "停车", "路线", "怎么去"]),
        ("历史文化", ["历史", "文化", "佛教", "玄奘", "祥符", "禅", "典故"]),
        ("游玩建议", ["路线", "推荐", "适合", "体验", "游览", "怎么玩", "打卡"]),
        ("便民服务", ["轮椅", "婴儿车", "寄存", "洗手间", "母婴", "WiFi"]),
        ("票务价格", ["成人票", "半价票", "免票", "联票", "票种"]),
        ("开放时间", ["开放时间", "营业时间", "闭园时间", "演出时间", "表演时间"]),
        ("景点介绍", ["景点", "建筑", "位置", "特色", "看点", "规模"]),
    ]
    for category, keywords in rules:
        if any(keyword in text for keyword in keywords):
            return category
    return "通用知识"


def dedupe_units(units: list) -> list:
    seen = set()
    result = []
    for unit in units:
        key = (normalize_key(unit.get("title", "")), normalize_key(unit.get("content", ""))[:120])
        if key in seen:
            continue
        seen.add(key)
        result.append(unit)
    for index, unit in enumerate(result, 1):
        unit["id"] = f"ku_{index:04d}"
    return result


def save_json(path: str, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"已保存: {path} ({os.path.getsize(path) / 1024:.1f} KB)")


def main():
    print("=" * 70)
    print("灵山胜境 raw Word 数据抽取")
    print("=" * 70)
    ensure_output_dir()
    require_files()

    print("抽取景点结构化表...")
    attractions = extract_attraction_tables(DOC_ATTRACTIONS)
    print(f"景点数量: {len(attractions)}")

    print("抽取指南补充表...")
    guide_tables = extract_guide_tables(DOC_GUIDE)
    print(f"指南表格数量: {len(guide_tables)}")
    attractions = merge_supplements(attractions, guide_tables)

    print("生成原文知识块...")
    units = []
    units.extend(attraction_to_knowledge_units(attractions))
    units.extend(split_guide_paragraphs(DOC_GUIDE))
    units.extend(guide_tables_to_knowledge_units(guide_tables))
    units = dedupe_units(units)
    print(f"知识块数量: {len(units)}")

    summary = {
        "raw_files": [os.path.basename(DOC_ATTRACTIONS), os.path.basename(DOC_GUIDE)],
        "attraction_count": len(attractions),
        "knowledge_unit_count": len(units),
        "categories": dict(sorted(category_counts(units).items())),
    }

    save_json(OUTPUT_ATTRACTIONS, attractions)
    save_json(OUTPUT_KNOWLEDGE_UNITS, units)
    save_json(OUTPUT_SOURCE_SUMMARY, summary)

    print("抽取完成。")


def category_counts(units: list) -> dict:
    counts = defaultdict(int)
    for unit in units:
        counts[unit.get("category", "通用知识")] += 1
    return counts


if __name__ == "__main__":
    main()
